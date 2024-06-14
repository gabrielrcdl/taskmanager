from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_abnt_document():
    # Cria um novo documento
    doc = Document()

    # Formatação ABNT
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Título
    title = doc.add_heading(level=1)
    title_run = title.add_run("Relatório de Testes de Software")
    title_run.bold = True
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Introdução
    doc.add_heading("Introdução", level=2)
    doc.add_paragraph("Este documento apresenta os resultados dos testes realizados no sistema de gerenciamento de tarefas.")

    # Resultados dos Testes
    doc.add_heading("Resultados dos Testes", level=2)
    test_results = [
        ("test_create_task", "Testa a criação de uma nova tarefa.", "Passou"),
        ("test_create_task_function", "Testa a função create_task do TaskManager.", "Passou"),
        ("test_create_task_with_empty_name", "Testa o comportamento da função create_task quando o nome da tarefa é uma string vazia.", "Falhou"),
        ("test_delete_task", "Testa a exclusão de uma tarefa existente.", "Passou"),
        ("test_delete_task_function", "Testa a função delete_task do TaskManager.", "Passou"),
        ("test_empty_task_list", "Testa se a lista de tarefas está vazia após a criação do objeto TaskManager.", "Passou"),
        ("test_update_task", "Testa a atualização de uma tarefa existente.", "Passou"),
        ("test_update_task_function", "Testa a função update_task do TaskManager.", "Passou")
    ]
    for result in test_results:
        test_name, test_description, test_status = result
        doc.add_paragraph(f"Teste: {test_name} - {test_status}")
        doc.add_paragraph(f"Descrição: {test_description}\n")

    # Salva o documento
    doc.save("relatorio_testes_abnt.docx")

if __name__ == "__main__":
    create_abnt_document()
